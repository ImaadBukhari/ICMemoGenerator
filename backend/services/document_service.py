import os
from typing import Dict, Any, List, Optional, Tuple
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import tempfile
import re

from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


from backend.db.models import MemoRequest, MemoSection

def create_memo_styles(doc: Document):
    """Create custom styles for the memo document"""
    
    styles = doc.styles
    
    # Choose your font here - change this to whatever you want
    DOCUMENT_FONT = 'Bangla Sangam MN'  
    
    # Title style
    if 'Memo Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Memo Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = DOCUMENT_FONT
        title_font.size = Pt(14)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(14)
    
    # Company name style
    if 'Company Name' not in [s.name for s in styles]:
        company_style = styles.add_style('Company Name', WD_STYLE_TYPE.PARAGRAPH)
        company_font = company_style.font
        company_font.name = DOCUMENT_FONT
        company_font.size = Pt(12)
        company_font.bold = True
        company_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_style.paragraph_format.space_after = Pt(12)
    
    # Section heading style
    if 'Section Heading' not in [s.name for s in styles]:
        heading_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = DOCUMENT_FONT
        heading_font.size = Pt(12)
        heading_font.bold = True
        heading_font.color.rgb = None 
        heading_style.paragraph_format.space_before = Pt(18)  
        heading_style.paragraph_format.space_after = Pt(12)   
    
    # Subsection heading style
    if 'Subsection Heading' not in [s.name for s in styles]:
        sub_heading_style = styles.add_style('Subsection Heading', WD_STYLE_TYPE.PARAGRAPH)
        sub_heading_font = sub_heading_style.font
        sub_heading_font.name = DOCUMENT_FONT
        sub_heading_font.size = Pt(10)
        sub_heading_font.bold = True
        sub_heading_style.paragraph_format.space_before = Pt(12)
        sub_heading_style.paragraph_format.space_after = Pt(8)
    
    # Assessment table heading style 
    if 'Assessment Table Heading' not in [s.name for s in styles]:
        assessment_table_style = styles.add_style('Assessment Table Heading', WD_STYLE_TYPE.PARAGRAPH)
        assessment_table_font = assessment_table_style.font
        assessment_table_font.name = DOCUMENT_FONT
        assessment_table_font.size = Pt(12)  
        assessment_table_font.bold = True
        assessment_table_style.paragraph_format.space_before = Pt(24) 
        assessment_table_style.paragraph_format.space_after = Pt(12)  
    
    # Body text style
    if 'Memo Body' not in [s.name for s in styles]:
        body_style = styles.add_style('Memo Body', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = DOCUMENT_FONT
        body_font.size = Pt(10)
        body_style.paragraph_format.line_spacing = 1.2  
        body_style.paragraph_format.space_after = Pt(8) 
    
    # List Bullet style 
    try:
        list_bullet_style = styles['List Bullet']
        list_bullet_font = list_bullet_style.font
        list_bullet_font.name = DOCUMENT_FONT
        list_bullet_font.size = Pt(10)
        list_bullet_style.paragraph_format.space_after = Pt(4) 
    except KeyError:
        # Create List Bullet style if it doesn't exist
        list_bullet_style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        list_bullet_font = list_bullet_style.font
        list_bullet_font.name = DOCUMENT_FONT
        list_bullet_font.size = Pt(10)
        list_bullet_style.paragraph_format.left_indent = Inches(0.25)
        list_bullet_style.paragraph_format.space_after = Pt(4)  
        list_bullet_style.paragraph_format.line_spacing = 1.2 

def clean_markdown_formatting(content: str, section_name: str) -> str:
    """
    Clean markdown formatting from content based on section type
    """
    # For executive summary and company snapshot, remove ## and ### headers entirely
    # This is because these sections should be concise summaries without subheaders
    if section_name in ['executive_summary', 'company_snapshot']:
        # Remove lines that start with ##, ###, or ####
        lines = content.split('\n')
        cleaned_lines = []
        for line in lines:
            stripped_line = line.strip()
            if not (stripped_line.startswith('####') or stripped_line.startswith('###') or stripped_line.startswith('##')):
                cleaned_lines.append(line)
        content = '\n'.join(cleaned_lines)
    
    # Remove extra whitespace and normalize line breaks
    content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)  # Max 2 consecutive line breaks
    content = content.strip()
    
    return content

def parse_formatted_content(content: str, section_name: str) -> List[Dict[str, Any]]:
    """
    Simple parser: Split content into lines and process basic formatting.
    - Lines starting with #, ##, ###, #### -> bold header (10pt, bold)
    - Lines with **text** -> paragraph with inline bold
    - Everything else -> regular paragraph (10pt, not bold)
    """
    import re
    content_blocks = []
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            # Empty line - add spacing
            content_blocks.append({'type': 'paragraph', 'content': ''})
        elif stripped.startswith('#'):
            # Hash pattern at line start -> bold header (10pt, bold)
            text = re.sub(r'^#+\s*', '', stripped)
            if text:
                content_blocks.append({
                    'type': 'bold_header',
                    'content': text
                })
        else:
            # Regular paragraph - check for inline **text** patterns
            content_blocks.append({
                'type': 'paragraph',
                'content': line
            })
    
    return content_blocks

def add_formatted_text_to_paragraph(paragraph, text: str, font_name: str = 'Bangla Sangam MN', font_size: int = 10):
    """
    Add text to a paragraph with proper bold formatting for **text** and consistent font
    """
    # Split text by bold markers
    parts = re.split(r'\*\*([^*]+)\*\*', text)
    
    for i, part in enumerate(parts):
        if not part:  # Skip empty parts
            continue
            
        if i % 2 == 0:  # Regular text (even indices)
            run = paragraph.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(font_size)
        else:  # Bold text (odd indices, content between **)
            run = paragraph.add_run(part)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size)

def extract_rating_from_content(content: str) -> Tuple[Optional[str], str]:
    """
    Extract rating number and clean description from assessment content
    Returns: (rating, cleaned_content)
    """
    # Look for rating patterns like "Rating: 7/10", "Score: 8", "7 out of 10", etc.
    rating_patterns = [
        r'rating[:\s]+(\d+)(?:/10|\s*out\s*of\s*10)?',
        r'score[:\s]+(\d+)(?:/10|\s*out\s*of\s*10)?',
        r'(\d+)/10',
        r'(\d+)\s*out\s*of\s*10',
        r'rating[:\s]+(\d+)',
        r'score[:\s]+(\d+)'
    ]
    
    rating = None
    for pattern in rating_patterns:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            rating = f"{match.group(1)}/10"
            break
    
    # Clean the content - remove redundant rating statements
    cleaned_content = content
    for pattern in rating_patterns:
        cleaned_content = re.sub(pattern, '', cleaned_content, flags=re.IGNORECASE)
    
    # Clean up extra whitespace and formatting
    cleaned_content = re.sub(r'\n\s*\n', '\n\n', cleaned_content)
    cleaned_content = cleaned_content.strip()
    
    return rating, cleaned_content

def create_assessment_table(doc: Document, assessment_sections: Dict[str, MemoSection]):
    """Create a professional assessment summary table with improved spacing"""
    
    # Set the font for table content - consistent with document
    TABLE_FONT = 'Bangla Sangam MN'
    
    # Add spacing before table
    doc.add_paragraph()  # Extra spacing above table
    
    # Add table heading with better spacing
    table_heading = doc.add_paragraph("Investment Assessment Summary", style='Assessment Table Heading')
    table_heading.paragraph_format.space_after = Pt(12)  # More space after heading
    
    # Create table with 3 columns: Category, Rating, Key Points
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Set column widths
    table.columns[0].width = Inches(2.2)  # Category - slightly wider
    table.columns[1].width = Inches(1.0)  # Rating  
    table.columns[2].width = Inches(4.3)  # Key Points - slightly narrower for better balance
    
    # Set table margins and spacing
    for row in table.rows:
        row.height = Inches(0.4)  # Minimum row height for better spacing
        
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Assessment Category"
    header_cells[1].text = "Rating"
    header_cells[2].text = "Key Points"
    
    # Format header row with gray background and white text
    for cell in header_cells:
        # Set background color using proper XML method
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), "434343")  # Gray background
        tc_pr.append(shd)
        
        # Add cell margins for better spacing
        tc_mar = OxmlElement('w:tcMar')
        
        # Set margins: top, left, bottom, right (in twentieths of a point)
        for margin_name, margin_value in [('top', '120'), ('left', '120'), ('bottom', '120'), ('right', '120')]:
            margin_elem = OxmlElement(f'w:{margin_name}')
            margin_elem.set(qn('w:w'), margin_value)
            margin_elem.set(qn('w:type'), 'dxa')
            tc_mar.append(margin_elem)
        
        tc_pr.append(tc_mar)
        
        # Format text - white and bold
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)  # Remove extra paragraph spacing
            for run in paragraph.runs:
                run.font.name = TABLE_FONT
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Assessment categories mapping
    assessment_mapping = {
        "assessment_people": "Team & Leadership",
        "assessment_market_opportunity": "Market Opportunity", 
        "assessment_product": "Product & Technology",
        "assessment_financials": "Financial Health",
        "assessment_traction_validation": "Traction & Validation",
        "assessment_deal_considerations": "Deal Structure"
    }
    
    # Add rows for each assessment
    for section_key, section_title in assessment_mapping.items():
        if section_key in assessment_sections:
            section = assessment_sections[section_key]
            
            # Extract rating and clean content
            rating, cleaned_content = extract_rating_from_content(section.content)
            
            # Add row to table
            row_cells = table.add_row().cells
            
            # Set minimum row height for better spacing
            table.rows[-1].height = Inches(0.6)  # Taller rows for content
            
            # Category name (first column) with light green background
            row_cells[0].text = section_title
            
            # Set light green background for category cell
            tc_pr = row_cells[0]._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), "a6ddce")  # Light green background
            tc_pr.append(shd)
            
            # Add cell margins for category cell
            tc_mar = OxmlElement('w:tcMar')
            for margin_name, margin_value in [('top', '120'), ('left', '120'), ('bottom', '120'), ('right', '120')]:
                margin_elem = OxmlElement(f'w:{margin_name}')
                margin_elem.set(qn('w:w'), margin_value)
                margin_elem.set(qn('w:type'), 'dxa')
                tc_mar.append(margin_elem)
            tc_pr.append(tc_mar)
            
            # Format category cell text
            for paragraph in row_cells[0].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = TABLE_FONT
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            
            # Rating (second column) - add margins
            rating_text = rating if rating else "N/A"
            row_cells[1].text = rating_text
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add margins to rating cell
            tc_pr_rating = row_cells[1]._tc.get_or_add_tcPr()
            tc_mar_rating = OxmlElement('w:tcMar')
            for margin_name, margin_value in [('top', '120'), ('left', '120'), ('bottom', '120'), ('right', '120')]:
                margin_elem = OxmlElement(f'w:{margin_name}')
                margin_elem.set(qn('w:w'), margin_value)
                margin_elem.set(qn('w:type'), 'dxa')
                tc_mar_rating.append(margin_elem)
            tc_pr_rating.append(tc_mar_rating)
            
            # Format rating cell
            for paragraph in row_cells[1].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = TABLE_FONT
                    run.font.bold = True
                    run.font.size = Pt(10)
            
            # Key points (third column) - add margins and better bullet spacing
            key_points = summarize_assessment_content_with_spacing(cleaned_content)
            row_cells[2].text = key_points
            
            # Add margins to key points cell
            tc_pr_points = row_cells[2]._tc.get_or_add_tcPr()
            tc_mar_points = OxmlElement('w:tcMar')
            for margin_name, margin_value in [('top', '120'), ('left', '120'), ('bottom', '120'), ('right', '120')]:
                margin_elem = OxmlElement(f'w:{margin_name}')
                margin_elem.set(qn('w:w'), margin_value)
                margin_elem.set(qn('w:type'), 'dxa')
                tc_mar_points.append(margin_elem)
            tc_pr_points.append(tc_mar_points)
            
            # Format key points cell
            for paragraph in row_cells[2].paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)  # Small spacing between bullet lines
                paragraph.paragraph_format.line_spacing = 1.2  # Better line spacing
                for run in paragraph.runs:
                    run.font.name = TABLE_FONT
                    run.font.size = Pt(10)
    
    # Add spacing after table
    doc.add_paragraph()  # Extra spacing below table
    doc.add_paragraph()  # Even more spacing before next content

def summarize_assessment_content_with_spacing(content: str) -> str:
    """Summarize assessment content with better bullet point spacing"""
    
    # Clean the content first
    cleaned_content = clean_table_content(content)
    
    # Split content into sentences/points
    sentences = re.split(r'[.!?]+', cleaned_content)
    sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]
    
    # Take first few key sentences and format as bullet points
    key_points = []
    for sentence in sentences[:3]:  # Take up to 3 key points
        if len(sentence) > 20:  # Only meaningful sentences
            sentence = sentence.strip()
            if not sentence.endswith('.'):
                sentence += '.'
            key_points.append(f"• {sentence}")
    
    # If we have fewer than 2 points, try to extract from original bullet points
    if len(key_points) < 2:
        # Look for sentences that might be key points
        lines = cleaned_content.split('\n')
        for line in lines:
            clean_line = line.strip()
            if len(clean_line) > 15 and clean_line not in [kp[2:] for kp in key_points]:
                if not clean_line.endswith('.'):
                    clean_line += '.'
                key_points.append(f"• {clean_line}")
                if len(key_points) >= 3:
                    break
    
    # Join with line breaks for better spacing between bullets
    return '\n\n'.join(key_points[:3]) if key_points else cleaned_content[:150] + "..." if len(cleaned_content) > 150 else cleaned_content

def clean_table_content(content: str) -> str:
    """
    Thoroughly clean content for table display - remove all markdown formatting
    """
    # Remove all hash markers (# ## ### #### etc.)
    content = re.sub(r'#+\s*', '', content)
    
    # Remove bold markers but keep the text
    content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
    
    # Remove bullet point markers at start of lines (we'll add our own)
    content = re.sub(r'^[•\-]\s*', '', content, flags=re.MULTILINE)
    
    # Clean up multiple spaces and line breaks
    content = re.sub(r'\s+', ' ', content)  # Multiple spaces to single space
    content = re.sub(r'\n\s*\n', '\n', content)  # Multiple line breaks to single
    
    # Remove any remaining markdown artifacts
    content = re.sub(r'[`~_]', '', content)  # Remove backticks, tildes, underscores
    
    return content.strip()


def add_header_footer(doc: Document, company_name: str, generation_date: str):
    """Add header and footer to the document"""
    
    HEADER_FOOTER_FONT = 'Bangla Sangam MN'  # Change this to match your preferred font
    
    # Add header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"CONFIDENTIAL - Investment Committee Memo"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Format header font
    for run in header_para.runs:
        run.font.name = HEADER_FOOTER_FONT
        run.font.size = Pt(8)
    
    # Add footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated on {generation_date} | Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Format footer font
    for run in footer_para.runs:
        run.font.name = HEADER_FOOTER_FONT
        run.font.size = Pt(8)

def process_markdown_bold(text: str) -> Tuple[str, List[Tuple[int, int]]]:
    """
    Process inline **text** patterns only.
    Correctly handles character offset when removing markers.
    
    Returns:
        (cleaned_text, bold_ranges) where bold_ranges is a list of (start, end) tuples
    """
    import re
    bold_ranges = []
    
    # Find all matches in original text
    matches = list(re.finditer(r'\*\*([^*]+)\*\*', text))
    
    if not matches:
        return text, []
    
    # Build cleaned text and track bold ranges
    # Process from start to end, building the cleaned text incrementally
    cleaned_text = ""
    current_pos = 0
    
    for match in matches:
        match_start = match.start()
        match_end = match.end()
        bold_text = match.group(1)
        
        # Add text before this match
        cleaned_text += text[current_pos:match_start]
        
        # Track bold range in cleaned_text
        bold_start = len(cleaned_text)
        bold_end = bold_start + len(bold_text)
        bold_ranges.append((bold_start, bold_end))
        
        # Add the bold text (without markers)
        cleaned_text += bold_text
        
        # Move past the match in original text
        current_pos = match_end
    
    # Add remaining text after last match
    cleaned_text += text[current_pos:]
    
    # Sort ranges by start position (should already be sorted)
    bold_ranges.sort(key=lambda x: x[0])
    
    return cleaned_text, bold_ranges

def format_section_content(content: str, section_name: str) -> List[Dict[str, Any]]:
    """
    Simple formatter: Process content into blocks with basic formatting.
    """
    content_blocks = parse_formatted_content(content, section_name)
    formatted_blocks = []
    
    for block in content_blocks:
        if block['type'] == 'paragraph':
            text = block['content']
            # Process inline **text** patterns for bold
            cleaned_text, bold_ranges = process_markdown_bold(text)
            formatted_blocks.append({
                'type': 'paragraph',
                'content': cleaned_text,
                'bold_ranges': bold_ranges
            })
        else:
            # Pass through other block types (bold_header, etc.)
            formatted_blocks.append(block)
    
    return formatted_blocks

def add_section_to_document(doc: Document, section_name: str, content: str, section_order: Dict[str, str]):
    """Add a section to the Word document with enhanced formatting"""
    
    BODY_FONT = 'Bangla Sangam MN'  # Consistent font for all body text
    BODY_SIZE = 10
    
    # Get section title
    section_title = section_order.get(section_name, section_name.replace('_', ' ').title())
    
    # Add section heading
    heading_para = doc.add_paragraph(section_title, style='Section Heading')
    
    # Format and add content with enhanced parsing
    formatted_blocks = format_section_content(content, section_name)
    
    for block in formatted_blocks:
        if block['type'] == 'subsection_header':
            # Add subsection header
            sub_para = doc.add_paragraph(block['content'], style='Subsection Heading')
        
        elif block['type'] == 'paragraph':
            # Add paragraph with bold formatting (#### becomes **bold** within paragraph)
            para = doc.add_paragraph(style='Memo Body')
            add_formatted_text_to_paragraph(para, block['content'], BODY_FONT, BODY_SIZE)
        
        elif block['type'] == 'bullet_list':
            # Add bullet list items with consistent formatting
            for item in block['items']:
                bullet_para = doc.add_paragraph(style='List Bullet')
                add_formatted_text_to_paragraph(bullet_para, item, BODY_FONT, BODY_SIZE)

def build_section_blocks(sections_dict: Dict[str, Any], assessment_sections: Dict[str, Any] = None) -> List[Dict[str, Any]]:
    """
    Build structured blocks from sections_dict for Google Docs creation.
    Returns a list of block dicts that can be used by create_google_doc_from_blocks.
    """
    blocks = []
    
    section_order = {
        "executive_summary": "Executive Summary",
        "company_snapshot": "Company Snapshot",
        "people": "Team & Leadership",
        "market_opportunity": "Market Opportunity", 
        "competitive_landscape": "Competitive Landscape",
        "product": "Product & Technology",
        "financial": "Financial Analysis",
        "traction_validation": "Traction & Validation",
        "deal_considerations": "Deal Considerations"
    }
    
    assessment_mapping = {
        "assessment_people": "Team & Leadership",
        "assessment_market_opportunity": "Market Opportunity", 
        "assessment_product": "Product & Technology",
        "assessment_financials": "Financial Health",
        "assessment_traction_validation": "Traction & Validation",
        "assessment_deal_considerations": "Deal Structure"
    }
    
    # Add Executive Summary first
    if "executive_summary" in sections_dict:
        section = sections_dict["executive_summary"]
        content = section.content if hasattr(section, 'content') else section
        blocks.append({
            'type': 'section_heading',
            'content': 'Executive Summary'
        })
        formatted_blocks = format_section_content(content, "executive_summary")
        blocks.extend(formatted_blocks)
        blocks.append({'type': 'paragraph', 'content': ''})  # Spacing
    
    # Add Assessment Summary as formatted text (not table)
    if assessment_sections:
        blocks.append({
            'type': 'section_heading',
            'content': 'Investment Assessment Summary'
        })
        
        # Format each assessment as readable text with category, rating, and full justification
        for section_key, section_title in assessment_mapping.items():
            if section_key in assessment_sections:
                section = assessment_sections[section_key]
                content = section.content if hasattr(section, 'content') else section
                
                # Extract rating and full content
                rating, cleaned_content = extract_rating_from_content(content)
                
                # Format as: "Category: Rating\nFull justification text"
                assessment_text = f"{section_title}"
                if rating:
                    assessment_text += f": {rating}"
                assessment_text += "\n"
                if cleaned_content:
                    assessment_text += cleaned_content
                
                # Category and rating as bold header (10pt)
                header_text = section_title
                if rating:
                    header_text += f": {rating}"
                blocks.append({
                    'type': 'bold_header',
                    'content': header_text
                })
                
                # Justification as paragraph
                if cleaned_content:
                    formatted_blocks = format_section_content(cleaned_content, section_key)
                    blocks.extend(formatted_blocks)
        
        blocks.append({'type': 'paragraph', 'content': ''})  # Spacing
    
    # Add main sections in order
    for section_key, section_title in section_order.items():
        if section_key in sections_dict and section_key != "executive_summary":
            section = sections_dict[section_key]
            content = section.content if hasattr(section, 'content') else section
            
            blocks.append({
                'type': 'section_heading',
                'content': section_title
            })
            
            formatted_blocks = format_section_content(content, section_key)
            blocks.extend(formatted_blocks)
            blocks.append({'type': 'paragraph', 'content': ''})  # Spacing
    
    return blocks

def add_sources_section(doc: Document, sections: List):
    """Add a sources/references section at the end of the document"""
    
    # Collect all unique sources from all sections
    all_sources = set()
    
    for section in sections:
        if section.data_sources and isinstance(section.data_sources, list):
            all_sources.update(section.data_sources)
    
    if not all_sources:
        return
    
    # Add page break before sources
    doc.add_page_break()
    
    # Add "Sources" heading
    sources_heading = doc.add_heading('Sources & References', level=1)
    sources_heading.style = 'Heading 1'
    
    # Add description
    desc = doc.add_paragraph()
    desc.add_run("The following sources were used in the preparation of this investment memo:").italic = True
    desc.paragraph_format.space_after = Pt(12)
    
    # Sort sources alphabetically and add them
    sorted_sources = sorted(list(all_sources))
    
    for i, source in enumerate(sorted_sources, 1):
        source_para = doc.add_paragraph(style='List Number')
        source_para.text = source
        
        # Format the source text
        for run in source_para.runs:
            run.font.name = 'Bangla Sangam MN'
            run.font.size = Pt(10)
        
        source_para.paragraph_format.space_after = Pt(6)
        source_para.paragraph_format.left_indent = Inches(0.25)

# Update the generate_word_document function to call this

def generate_google_doc(user, db: Session, sections_dict: Dict[str, Any], company_name: str) -> str:
    """
    Generate a Google Doc from memo sections and save it to Investments folder.
    
    Args:
        user: User object (for Drive access)
        db: Database session
        sections_dict: Dictionary of section_name -> MemoSection objects
        company_name: Company name for document title
        
    Returns:
        Google Doc URL
    """
    from backend.services.google_service import (
        create_google_doc_from_blocks
    )
    
    try:
        print(f"Starting Google Doc generation for {company_name}")
        
        # Separate assessment sections from main sections
        assessment_sections = {k: v for k, v in sections_dict.items() if k.startswith('assessment_')}
        main_sections_dict = {k: v for k, v in sections_dict.items() if not k.startswith('assessment_')}
        
        # Build blocks from sections
        blocks = build_section_blocks(main_sections_dict, assessment_sections)
        
        # Add sources section at the end
        all_sources = set()
        for section in sections_dict.values():
            if hasattr(section, 'data_sources') and section.data_sources:
                if isinstance(section.data_sources, list):
                    all_sources.update(section.data_sources)
        
        if all_sources:
            blocks.append({
                'type': 'heading',
                'content': 'Sources'
            })
            for source in sorted(all_sources):
                blocks.append({
                    'type': 'paragraph',
                    'content': source
                })
        
        # Create document title with date in DD/MM/YY format
        generation_date = datetime.now().strftime("%d/%m/%y")
        
        # Add title and subtitle blocks at the beginning
        header_blocks = [
            {
                'type': 'title',
                'content': f"{company_name} - {generation_date}"
            },
            {
                'type': 'subtitle',
                'content': 'Series A Memo'
            },
            {
                'type': 'subtitle',
                'content': 'Deal Team: [To Be Determined]'  # Placeholder for now
            },
            {
                'type': 'paragraph',
                'content': ''  # Spacing
            }
        ]
        
        # Prepend header blocks to content blocks
        blocks = header_blocks + blocks
        
        # Create document title for file name: IC Memo_[company name] DD_MM_YYYY
        generation_date_for_filename = datetime.now().strftime("%d_%m_%Y")
        doc_title = f"IC Memo_{company_name} {generation_date_for_filename}"
        
        # Create Google Doc (no folder, will be in user's Drive root)
        doc_url = create_google_doc_from_blocks(
            user=user,
            db=db,
            title=doc_title,
            blocks=blocks,
            parent_folder_id=None
        )
        
        print(f"✅ Google Doc created: {doc_url}")
        return doc_url
        
    except ValueError as e:
        if "No Google tokens found" in str(e):
            error_msg = "Google Drive connection required. Please connect your Google account via OAuth to generate documents."
            print(f"❌ Error generating Google Doc: {error_msg}")
            raise ValueError(error_msg)
        else:
            print(f"❌ Error generating Google Doc: {str(e)}")
            import traceback
            traceback.print_exc()
            raise
    except Exception as e:
        print(f"❌ Error generating Google Doc: {str(e)}")
        import traceback
        traceback.print_exc()
        raise

def generate_word_document(db: Session, memo_request_id: int) -> Optional[str]:
    """Generate a formatted Word document from memo sections with sources"""
    
    try:
        print(f"Starting document generation for memo {memo_request_id}")
        
        # Check if python-docx is available
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            print("✅ python-docx is available")
        except ImportError as e:
            print(f"❌ python-docx import error: {str(e)}")
            return None
        
        # Get memo request
        memo_request = db.query(MemoRequest).filter(
            MemoRequest.id == memo_request_id
        ).first()
        
        if not memo_request:
            print(f"❌ Memo request {memo_request_id} not found")
            return None
        
        print(f"✅ Found memo request for {memo_request.company_name}")
        
        # Get all completed sections
        sections = db.query(MemoSection).filter(
            MemoSection.memo_request_id == memo_request_id,
            MemoSection.status == "completed"
        ).order_by(MemoSection.created_at).all()
        
        print(f"Found {len(sections)} completed sections")
        
        if not sections:
            print("❌ No completed sections found for this memo")
            return None
        
        for section in sections:
            print(f"  - {section.section_name}: {len(section.content) if section.content else 0} chars")
        
        # Create document
        print("Creating Word document...")
        doc = Document()
        create_memo_styles(doc)
        
        # Add document title and company info
        title_para = doc.add_paragraph("INVESTMENT COMMITTEE MEMO", style='Memo Title')
        company_para = doc.add_paragraph(memo_request.company_name, style='Company Name')
        
        # Add generation info
        generation_date = datetime.now().strftime("%B %d, %Y")
        info_para = doc.add_paragraph(f"Generated: {generation_date}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in info_para.runs:
            run.font.name = 'Bangla Sangam MN'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()  # Spacing
        
        print("✅ Added title and header")
        
        # Add header and footer
        add_header_footer(doc, memo_request.company_name, generation_date)
        
        # Create sections lookup
        sections_dict = {section.section_name: section for section in sections}
        
        # Separate assessment sections from main sections
        assessment_sections = {k: v for k, v in sections_dict.items() if k.startswith('assessment_')}
        main_sections_dict = {k: v for k, v in sections_dict.items() if not k.startswith('assessment_')}
        
        print(f"Main sections: {list(main_sections_dict.keys())}")
        print(f"Assessment sections: {list(assessment_sections.keys())}")
        
        # Add Executive Summary first
        if "executive_summary" in main_sections_dict:
            print("Adding executive summary...")
            add_section_to_document(doc, "executive_summary", main_sections_dict["executive_summary"].content, 
                                  {"executive_summary": "Executive Summary"})
            doc.add_paragraph()  # Spacing
        
        # Add Assessment Summary Table
        if assessment_sections:
            print("Adding assessment table...")
            create_assessment_table(doc, assessment_sections)
            doc.add_page_break()
        
        # Main section order (excluding executive summary since we added it first)
        main_section_order = {
            "company_snapshot": "Company Snapshot",
            "people": "Team & Leadership",
            "market_opportunity": "Market Opportunity", 
            "competitive_landscape": "Competitive Landscape",
            "product": "Product & Technology",
            "financial": "Financial Analysis",
            "traction_validation": "Traction & Validation",
            "deal_considerations": "Deal Considerations"
        }
        
        # Add main sections in proper order with enhanced formatting
        for section_key, section_title in main_section_order.items():
            if section_key in main_sections_dict:
                print(f"Adding section with formatting: {section_title}")
                add_section_to_document(doc, section_key, main_sections_dict[section_key].content, 
                                      {section_key: section_title})
                doc.add_paragraph()  # Spacing between sections
        
        # ADD SOURCES SECTION HERE
        print("Adding sources section...")
        add_sources_section(doc, sections)
        
        # Ensure documents directory exists
        docs_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'generated_docs')
        print(f"Documents directory: {docs_dir}")
        
        os.makedirs(docs_dir, exist_ok=True)
        print("✅ Documents directory created/verified")
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_company_name = memo_request.company_name.replace(' ', '_').replace('/', '_')
        filename = f"IC_Memo_{safe_company_name}_{timestamp}.docx"
        file_path = os.path.join(docs_dir, filename)
        
        print(f"Saving document to: {file_path}")
        
        # Save the document
        doc.save(file_path)
        print("✅ Word document saved successfully with enhanced formatting")
        
        # Verify file exists
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            print(f"✅ Document file exists, size: {file_size} bytes")
            return file_path
        else:
            print("❌ Document file was not created")
            return None
        
    except Exception as e:
        print(f"❌ Error generating Word document: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def get_document_summary(db: Session, memo_request_id: int) -> Dict[str, Any]:
    """Get a summary of document generation status"""
    
    memo_request = db.query(MemoRequest).filter(
        MemoRequest.id == memo_request_id
    ).first()
    
    if not memo_request:
        return {"error": "Memo request not found"}
    
    sections = db.query(MemoSection).filter(
        MemoSection.memo_request_id == memo_request_id
    ).all()
    
    completed_sections = [s for s in sections if s.status == "completed"]
    failed_sections = [s for s in sections if s.status == "failed"]
    
    return {
        "memo_id": memo_request_id,
        "company_name": memo_request.company_name,
        "total_sections": len(sections),
        "completed_sections": len(completed_sections),
        "failed_sections": len(failed_sections),
        "success_rate": len(completed_sections) / len(sections) if sections else 0,
        "status": memo_request.status,
        "sections_detail": [
            {
                "name": s.section_name,
                "status": s.status,
                "content_length": len(s.content) if s.content else 0,
                "error": s.error_log if s.status == "failed" else None
            }
            for s in sections
        ]
    }


# --------------------------------------------------------------------------
# 1. Create styles
# --------------------------------------------------------------------------
def create_short_memo_styles(doc: Document):
    """Create consistent styles for the short memo document"""
    styles = doc.styles
    FONT = "Bangla Sangam MN"

    def ensure_style(name, style_type, config_fn):
        if name not in [s.name for s in styles]:
            style = styles.add_style(name, style_type)
            config_fn(style)

    ensure_style("Short Title", WD_STYLE_TYPE.PARAGRAPH, lambda s: (
        setattr(s.font, "name", FONT),
        setattr(s.font, "size", Pt(14)),
        setattr(s.font, "bold", True),
        setattr(s.paragraph_format, "alignment", WD_ALIGN_PARAGRAPH.CENTER),
        setattr(s.paragraph_format, "space_after", Pt(12))
    ))

    ensure_style("Short Subtitle", WD_STYLE_TYPE.PARAGRAPH, lambda s: (
        setattr(s.font, "name", FONT),
        setattr(s.font, "size", Pt(12)),
        setattr(s.font, "bold", True),
        setattr(s.paragraph_format, "alignment", WD_ALIGN_PARAGRAPH.CENTER),
        setattr(s.paragraph_format, "space_after", Pt(10))
    ))

    ensure_style("Short Body", WD_STYLE_TYPE.PARAGRAPH, lambda s: (
        setattr(s.font, "name", FONT),
        setattr(s.font, "size", Pt(10)),
        setattr(s.paragraph_format, "space_after", Pt(6)),
        setattr(s.paragraph_format, "line_spacing", 1.2)
    ))

    ensure_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH, lambda s: (
        setattr(s.font, "name", FONT),
        setattr(s.font, "size", Pt(10)),
        setattr(s.paragraph_format, "left_indent", Inches(0.25)),
        setattr(s.paragraph_format, "space_after", Pt(4))
    ))


# --------------------------------------------------------------------------
# 2. Add header
# --------------------------------------------------------------------------
def add_short_memo_header(doc: Document, company_name: str, series_text: str):
    """Add the company name and memo subtitle at the top"""
    doc.add_paragraph(company_name, style="Short Title")

    subtitle_para = doc.add_paragraph(f"Initial {series_text} IC Memo", style="Short Subtitle")
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generation_date = datetime.now().strftime("%B %d, %Y")
    date_para = doc.add_paragraph(f"Generated {generation_date}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        run.font.name = "Bangla Sangam MN"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(127, 140, 141)

    doc.add_paragraph()  # spacing


# --------------------------------------------------------------------------
# 3. Add Problem/Solution Table
# --------------------------------------------------------------------------
def add_short_problem_solution_table(doc: Document, problem_text: str, solution_text: str):
    """Add a formatted 2x2 Problem/Solution table"""
    FONT = "Bangla Sangam MN"

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Set column widths more explicitly
    table.columns[0].width = Inches(1.0)  # Header column
    table.columns[1].width = Inches(5.0)  # Content column
    
    # Ensure table width is set to match column widths
    table.width = Inches(6.0)  # Total width = 1.0 + 5.0
    
    # Force column widths using XML properties
    for i, col in enumerate(table.columns):
        col_width = Inches(1.0) if i == 0 else Inches(5.0)
        col.width = col_width
        
        # Set column width in XML
        for cell in [row.cells[i] for row in table.rows]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_width.inches * 1440)))  # Convert to twentieths of a point
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    headers = ["Problem:", "Solution:"]
    values = [problem_text or "Problem not provided", solution_text or "Solution not provided"]

    for i in range(2):
        # Header cell
        header_cell = table.rows[i].cells[0]
        header_cell.text = headers[i]

        # Apply shading (light green)
        tc_pr = header_cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "a6ddce")
        tc_pr.append(shd)

        for run in header_cell.paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(11)
            run.font.bold = True

        # Value cell
        value_cell = table.rows[i].cells[1]
        para = value_cell.paragraphs[0]
        add_formatted_text_to_paragraph(para, values[i], FONT, 10)
        para.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()  # spacing below table


# --------------------------------------------------------------------------
# 4. Format content parsing utilities (reuse from full memo)
# --------------------------------------------------------------------------
def clean_markdown_formatting_short(content: str) -> str:
    """Simplified markdown cleaner for short memo sections"""
    content = re.sub(r"#+\s*", "", content)  # remove headers
    content = re.sub(r"\*\*([^*]+)\*\*", r"\1", content)  # remove bold markers
    content = re.sub(r"\s+", " ", content)
    return content.strip()


def format_short_section_content(content: str) -> List[str]:
    """Break section content into bullet-style or paragraph chunks (for short memos)"""
    content = clean_markdown_formatting_short(content)
    lines = [l.strip() for l in content.split("\n") if l.strip()]
    return lines


def add_formatted_text_to_paragraph(paragraph, text: str, font_name: str, font_size: int):
    """Add text to a paragraph with inline bold parsing"""
    parts = re.split(r"\*\*([^*]+)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if i % 2 == 1:
            run.bold = True


# --------------------------------------------------------------------------
# 5. Add section blocks
# --------------------------------------------------------------------------
def add_short_section(doc: Document, section: MemoSection):
    """Add section content (without title) as bullet points or short paragraphs"""
    if not section or not section.content:
        return

    lines = format_short_section_content(section.content)
    for line in lines:
        # detect bullet or paragraph
        if line.startswith("- ") or line.startswith("• "):
            text = line[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            add_formatted_text_to_paragraph(para, f"• {text}", "Bangla Sangam MN", 10)
        else:
            para = doc.add_paragraph(line.strip(), style="Short Body")
            para.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()  # spacing after each section


# --------------------------------------------------------------------------
# 6. Add sources section
# --------------------------------------------------------------------------
def add_short_sources(doc: Document, sections: List[MemoSection]):
    """Add a list of unique sources at the end of the short memo"""
    all_sources = set()
    for section in sections:
        if section.data_sources:
            all_sources.update(section.data_sources)

    if not all_sources:
        return

    doc.add_paragraph("Sources", style="Short Subtitle")
    sorted_sources = sorted(list(all_sources))
    for idx, src in enumerate(sorted_sources, 1):
        para = doc.add_paragraph(f"[{idx}] {src}", style="Short Body")
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(4)


# --------------------------------------------------------------------------
# 7. Main function
# --------------------------------------------------------------------------
def generate_short_word_document(db: Session, memo_request_id: int) -> Optional[str]:
    """Generate a formatted short memo Word document"""
    try:
        print(f"Starting short memo generation for ID {memo_request_id}")

        memo_request = db.query(MemoRequest).filter(MemoRequest.id == memo_request_id).first()
        if not memo_request:
            print("❌ Memo request not found")
            return None

        sections = (
            db.query(MemoSection)
            .filter(MemoSection.memo_request_id == memo_request_id, MemoSection.status == "completed")
            .all()
        )
        if not sections:
            print("❌ No completed sections found")
            return None

        sections_dict = {s.section_name: s for s in sections}

        # Create document
        doc = Document()
        create_short_memo_styles(doc)

        # Header
        add_short_memo_header(doc, memo_request.company_name, "[X]")

        # Problem/Solution Table
        problem = sections_dict.get("problem")
        solution = sections_dict.get("solution")
        add_short_problem_solution_table(
            doc,
            problem.content if problem else "",
            solution.content if solution else ""
        )

        # Company Brief
        if "company_brief" in sections_dict:
            add_short_section(doc, sections_dict["company_brief"])

        # Bullet Point Sections
        ordered_sections = [
            "startup_overview",
            "founder_team",
            "deal_traction",
            "competitive_landscape",
            "remarks"
        ]

        for key in ordered_sections:
            section = sections_dict.get(key)
            if section:
                add_short_section(doc, section)

        # Sources
        add_short_sources(doc, sections)

        # Save document
        output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "generated_documents")
        os.makedirs(output_dir, exist_ok=True)
        filename = f"short_memo_{memo_request_id}_{memo_request.company_name.replace(' ', '_')}.docx"
        file_path = os.path.join(output_dir, filename)
        doc.save(file_path)
        print(f"✅ Short memo saved: {file_path}")
        return file_path

    except Exception as e:
        print(f"❌ Error generating short memo: {e}")
        import traceback
        traceback.print_exc()
        return None